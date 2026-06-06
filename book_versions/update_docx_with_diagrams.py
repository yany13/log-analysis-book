#!/usr/bin/env python3
"""
Update Chapters Overview DOCX with all diagrams mapped to correct chapters
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path

# Paths
DIAGRAM_DIR = r"D:\workspace\ai\code\log-analysis-book\diagrams"
OUTPUT_DOCX = r"D:\workspace\ai\code\log-analysis-book\book_versions\Chapters_Overviewv3.1.docx"

# Diagram mapping: chapter -> [(diagram_file, description)]
DIAGRAM_MAP = {
    "chapter_1": [
        ("01-infrastructure-architecture.svg", "RAG Pipeline Architecture"),
    ],
    "chapter_2": [
        ("02-application-flow.svg", "ByteBite Application Flow"),
        ("04-servlet-architecture.svg", "Servlet Architecture"),
        ("03-database-schema.svg", "Database Schema Overview"),
    ],
    "chapter_3": [
        ("08-docker-compose-hierarchy.svg", "Docker Compose Infrastructure"),
        ("01-infrastructure-architecture.svg", "AI Toolkit Architecture"),
    ],
    "chapter_4": [
        ("05-servlet-architecture-01-login-ldap-timeout.svg", "LDAP Timeout Failure Flow"),
    ],
    "chapter_5": [
        ("05-servlet-architecture-02-menu-db-leak.svg", "Database Connection Leak Scenario"),
    ],
    "chapter_6": [
        ("05-servlet-architecture-03-analytics-oom.svg", "OOM Memory Exhaustion"),
    ],
    "chapter_7": [
        ("05-servlet-architecture-04-kds-deadlock.svg", "Deadlock Lock Contention"),
    ],
    "chapter_8": [
        ("05-servlet-architecture-05-checkout-ssl.svg", "SSL Certificate Validation Failure"),
    ],
    "appendix_1": [
        ("06-database-schema-er-diagram.svg", "Complete ER Diagram"),
    ],
    "appendix_2": [
        ("app-project-folder-structure.svg", "Project Directory Structure"),
    ],
}

# Chapter content
CHAPTERS = {
    "chapter_1": {
        "title": "Chapter 1: Foundations - Intelligent Log Analysis",
        "description": "Establishing the rationale for leveraging AI in middleware troubleshooting, focusing on skill gaps, creating knowledge base, and introducing the architecture of our AI-powered diagnostic tools. We'll explore the core principles of RAG (Retrieval-Augmented Generation) and how it's optimized for enterprise operational data.",
    },
    "chapter_2": {
        "title": "Chapter 2: ByteBite Restaurant Application - Overview",
        "description": "A deep dive into custom application infrastructure architecture overview with underlying end-to-end application flow details.",
    },
    "chapter_3": {
        "title": "Chapter 3: The Log Analyzer AI Tool – Real-Time Log Investigation",
        "description": "A detailed exploration of the AI-powered tool, including its architecture, environment setup, and the utilization of Ollama and LangChain for secure LLM integration. Focuses on securing the data flow and anonymization techniques.",
    },
    "chapter_4": {
        "title": "Chapter 4: User Login Issues – Addressing LDAP Authentication Issues",
        "description": "Diagnosing authentication failures and timeouts. Detailed walkthroughs showcasing the issue production screenshots and resolution approaches using both traditional methods and AI-powered tools with multiple LLMs.",
        "curl_example": 'curl -X POST http://localhost:8080/bytebite/login -d "username=admin&password=admin"',
    },
    "chapter_5": {
        "title": "Chapter 5: Database Connection Pool Issues – Optimizing HikariCP Pool Utilization",
        "description": "Analyzing and resolving database connection leak scenarios using HikariCP. Focuses on how to identify and correct pool exhaustion issues.",
        "curl_example": "for i in {1..30}; do curl http://localhost:8080/bytebite/menu/api; done",
    },
    "chapter_6": {
        "title": "Chapter 6: Memory Management – Identifying JVM OOM",
        "description": "Analyzing memory-related issues in middleware environments. We explore the identification of JVM Out of Memory (OOM) errors and thread leaks.",
        "curl_example": 'curl "http://localhost:8080/bytebite/analytics/api?action=all-time-report"',
    },
    "chapter_7": {
        "title": "Chapter 7: Thread Analysis – Detecting and Resolving Thread Issues",
        "description": "Investigating thread-related problems, including blocked threads, thread pool exhaustion, and zombie threads.",
        "curl_example": 'curl -X POST http://localhost:8080/bytebite/kds -H "Content-Type: application/json" -d \'{"orderId":"ORD1","status":"IN"}\'',
    },
    "chapter_8": {
        "title": "Chapter 8: SSL Issues - Tracing Issues",
        "description": "Understanding SSL-specific use cases of handshake and truststore failure, how to utilize Wireshark to pinpoint handshaking issues, etc.",
        "curl_example": 'curl -X POST http://localhost:8080/bytebite/checkout/api -G --data-urlencode "injectSSLFailure=true" -H "Content-Type: application/json" -d \'{"totalAmount":99}\'',
    },
    "appendix_1": {
        "title": "Appendix 1: Database Schema Details",
        "description": "Complete database schema and ER diagram for the ByteBite application.",
    },
    "appendix_2": {
        "title": "Appendix 2: Project Structure",
        "description": "Complete project folder and file organization for ByteBite application.",
    },
}


def add_heading(doc, text, level=1):
    """Add heading to document"""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_paragraph_with_style(doc, text, style=None, bold=False, italic=False):
    """Add styled paragraph"""
    p = doc.add_paragraph(text, style=style)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p


def add_image_to_doc(doc, image_path, width=Inches(5.5), caption=None):
    """Add image with caption or reference if SVG"""
    if not os.path.exists(image_path):
        p = doc.add_paragraph("[Diagram: {}]".format(os.path.basename(image_path)))
        p.paragraph_format.left_indent = Inches(0.25)
        return

    # For SVG files, add a reference instead
    if image_path.endswith('.svg'):
        p = doc.add_paragraph()
        filename = os.path.basename(image_path)
        p.add_run("[Diagram: ").italic = False
        run = p.add_run(filename)
        run.italic = True
        p.add_run("]")

        if caption:
            caption_p = doc.add_paragraph(caption)
            caption_p.paragraph_format.left_indent = Inches(0.25)
            for run in caption_p.runs:
                run.font.size = Pt(9)
                run.font.italic = True
        return

    # For other image types, try to add directly
    try:
        p = doc.add_picture(image_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            caption_p = doc.add_paragraph(caption)
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption_p.runs:
                run.font.size = Pt(9)
                run.font.italic = True
    except Exception as e:
        p = doc.add_paragraph("[Error: {}]".format(str(e)))


def create_updated_docx():
    """Create updated DOCX with diagrams"""
    doc = Document()

    # Title page
    title = doc.add_heading("Middleware Troubleshooting Use Cases", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("A Traditional Approach with AI Assist", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Learn to design and deploy AI Agents that diagnose production failures through log analysis.")

    doc.add_page_break()

    # Table of Contents
    toc_heading = doc.add_heading("Table of Contents", level=1)

    for ch_key in ["chapter_1", "chapter_2", "chapter_3", "chapter_4", "chapter_5", "chapter_6", "chapter_7", "chapter_8"]:
        ch = CHAPTERS[ch_key]
        doc.add_paragraph(ch["title"], style="List Bullet")

    for app_key in ["appendix_1", "appendix_2"]:
        app = CHAPTERS[app_key]
        doc.add_paragraph(app["title"], style="List Bullet")

    doc.add_page_break()

    # Add chapters
    for ch_key in ["chapter_1", "chapter_2", "chapter_3", "chapter_4", "chapter_5", "chapter_6", "chapter_7", "chapter_8"]:
        ch = CHAPTERS[ch_key]
        add_heading(doc, ch["title"], level=1)
        add_paragraph_with_style(doc, ch["description"])

        # Add diagrams for this chapter
        if ch_key in DIAGRAM_MAP:
            doc.add_paragraph()  # Spacing
            for diagram_file, diagram_caption in DIAGRAM_MAP[ch_key]:
                diagram_path = os.path.join(DIAGRAM_DIR, diagram_file)
                add_image_to_doc(doc, diagram_path, width=Inches(5.5), caption=diagram_caption)
                doc.add_paragraph()  # Spacing

        # Add curl example if present
        if "curl_example" in ch:
            doc.add_heading("Example Command:", level=3)
            code_p = doc.add_paragraph(ch["curl_example"])
            for run in code_p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        doc.add_page_break()

    # Add appendices
    for app_key in ["appendix_1", "appendix_2"]:
        app = CHAPTERS[app_key]
        add_heading(doc, app["title"], level=1)
        add_paragraph_with_style(doc, app["description"])

        # Add diagrams for appendix
        if app_key in DIAGRAM_MAP:
            doc.add_paragraph()
            for diagram_file, diagram_caption in DIAGRAM_MAP[app_key]:
                diagram_path = os.path.join(DIAGRAM_DIR, diagram_file)
                add_image_to_doc(doc, diagram_path, width=Inches(5.5), caption=diagram_caption)
                doc.add_paragraph()

        doc.add_page_break()

    # Save document
    doc.save(OUTPUT_DOCX)
    print("[OK] Created: " + OUTPUT_DOCX)
    print("[OK] Document size: {:.2f} MB".format(os.path.getsize(OUTPUT_DOCX) / (1024*1024)))


if __name__ == "__main__":
    try:
        create_updated_docx()
        print("\n[SUCCESS] DOCX updated successfully with all diagrams!")
    except Exception as e:
        print("[ERROR] " + str(e))
        import traceback
        traceback.print_exc()
